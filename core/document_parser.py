"""
Multi-format Document Parser with Arabic support
Supports: PDF, DOCX, TXT
"""
import logging
from pathlib import Path
from typing import Optional, Dict, List
from abc import ABC, abstractmethod

import fitz  # PyMuPDF
import pdfplumber
import docx

from .arabic_processor import ArabicTextProcessor

logger = logging.getLogger(__name__)


class BaseParser(ABC):
    """Abstract base class for document parsers"""
    
    @abstractmethod
    def parse(self, file_path: Path) -> Dict[str, any]:
        """Parse document and return structured content"""
        pass


class PDFParser(BaseParser):
    """
    PDF Parser with dual-engine support
    Uses both PyMuPDF and pdfplumber for robust Arabic extraction
    """
    
    def __init__(self, arabic_processor: Optional[ArabicTextProcessor] = None):
        """
        Initialize PDF parser
        
        Args:
            arabic_processor: Optional ArabicTextProcessor instance
        """
        self.arabic_processor = arabic_processor or ArabicTextProcessor()
    
    def parse_with_pymupdf(self, file_path: Path) -> str:
        """
        Parse PDF using PyMuPDF (fitz)
        Generally better for text extraction with proper unicode handling
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text = ""
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                # Extract text with proper unicode handling
                page_text = page.get_text("text", sort=True)
                
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PyMuPDF parsing failed for {file_path}: {e}")
            raise
        
        return text
    
    def parse_with_pdfplumber(self, file_path: Path) -> str:
        """
        Parse PDF using pdfplumber
        Better for complex layouts and tables
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text = ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                        
        except Exception as e:
            logger.error(f"pdfplumber parsing failed for {file_path}: {e}")
            raise
        
        return text
    
    def parse(self, file_path: Path, engine: str = 'pdfplumber') -> Dict[str, any]:
        """
        Parse PDF file using specified engine
        
        Args:
            file_path: Path to PDF file
            engine: 'pymupdf', 'pdfplumber', or 'both'
            
        Returns:
            Dictionary with parsed content and metadata
        """
        logger.info(f"Parsing PDF: {file_path} with engine: {engine}")
        
        raw_text = ""
        
        if engine == 'pymupdf':
            raw_text = self.parse_with_pymupdf(file_path)
        elif engine == 'pdfplumber':
            raw_text = self.parse_with_pdfplumber(file_path)
        elif engine == 'both':
            # Try PyMuPDF first, fallback to pdfplumber if needed
            try:
                raw_text = self.parse_with_pymupdf(file_path)
            except Exception:
                logger.warning("PyMuPDF failed, trying pdfplumber...")
                raw_text = self.parse_with_pdfplumber(file_path)
        else:
            raise ValueError(f"Invalid engine: {engine}")
        
        # Process Arabic text
        processed = self.arabic_processor.process_text(raw_text, mode='dual')
        
        return {
            'raw_text': raw_text,
            'retrieval_text': processed.get('retrieval', ''),
            'search_text': processed.get('search', ''),
            'is_arabic': processed.get('is_arabic', False),
            'entities': processed.get('entities', []),
            'file_path': str(file_path),
            'file_type': 'pdf',
        }


class DOCXParser(BaseParser):
    """DOCX Parser with Arabic support"""
    
    def __init__(self, arabic_processor: Optional[ArabicTextProcessor] = None):
        """
        Initialize DOCX parser
        
        Args:
            arabic_processor: Optional ArabicTextProcessor instance
        """
        self.arabic_processor = arabic_processor or ArabicTextProcessor()
    
    def parse(self, file_path: Path) -> Dict[str, any]:
        """
        Parse DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with parsed content and metadata
        """
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            raw_text = "\n\n".join(paragraphs)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)
            
            if tables_text:
                raw_text += "\n\n--- Tables ---\n\n" + "\n".join(tables_text)
            
            # Process Arabic text
            processed = self.arabic_processor.process_text(raw_text, mode='dual')
            
            return {
                'raw_text': raw_text,
                'retrieval_text': processed.get('retrieval', ''),
                'search_text': processed.get('search', ''),
                'is_arabic': processed.get('is_arabic', False),
                'entities': processed.get('entities', []),
                'file_path': str(file_path),
                'file_type': 'docx',
                'num_paragraphs': len(paragraphs),
                'num_tables': len(doc.tables),
            }
            
        except Exception as e:
            logger.error(f"DOCX parsing failed for {file_path}: {e}")
            raise


class TXTParser(BaseParser):
    """Plain text parser with encoding detection"""
    
    def __init__(self, arabic_processor: Optional[ArabicTextProcessor] = None):
        """
        Initialize TXT parser
        
        Args:
            arabic_processor: Optional ArabicTextProcessor instance
        """
        self.arabic_processor = arabic_processor or ArabicTextProcessor()
    
    def parse(self, file_path: Path) -> Dict[str, any]:
        """
        Parse TXT file with encoding detection
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Dictionary with parsed content and metadata
        """
        logger.info(f"Parsing TXT: {file_path}")
        
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'utf-16', 'cp1256', 'iso-8859-6']
        
        raw_text = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    raw_text = f.read()
                used_encoding = encoding
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if raw_text is None:
            raise ValueError(f"Could not decode {file_path} with any known encoding")
        
        logger.info(f"Successfully decoded with encoding: {used_encoding}")
        
        # For TXT files, text is already in correct order (no RTL extraction needed)
        # Only apply normalization and entity extraction
        retrieval_text = self.arabic_processor.clean_whitespace(raw_text)
        search_text = self.arabic_processor.normalize_for_search(retrieval_text)
        is_arabic = self.arabic_processor.is_arabic_text(retrieval_text)
        entities = self.arabic_processor.extract_arabic_entities(retrieval_text) if is_arabic else []
        
        return {
            'raw_text': raw_text,
            'retrieval_text': retrieval_text,
            'search_text': search_text,
            'is_arabic': is_arabic,
            'entities': entities,
            'file_path': str(file_path),
            'file_type': 'txt',
            'encoding': used_encoding,
        }


class DocumentParser:
    """
    Unified document parser supporting multiple formats
    Automatically selects appropriate parser based on file extension
    """
    
    def __init__(self, arabic_processor: Optional[ArabicTextProcessor] = None):
        """
        Initialize document parser
        
        Args:
            arabic_processor: Optional ArabicTextProcessor instance
        """
        self.arabic_processor = arabic_processor or ArabicTextProcessor()
        
        # Initialize format-specific parsers
        self.parsers = {
            '.pdf': PDFParser(self.arabic_processor),
            '.docx': DOCXParser(self.arabic_processor),
            '.doc': DOCXParser(self.arabic_processor),
            '.txt': TXTParser(self.arabic_processor),
        }
    
    def parse(self, file_path: str | Path, **kwargs) -> Dict[str, any]:
        """
        Parse document of any supported format
        
        Args:
            file_path: Path to document
            **kwargs: Additional arguments passed to specific parser
            
        Returns:
            Dictionary with parsed content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        ext = file_path.suffix.lower()
        
        # Select parser
        parser = self.parsers.get(ext)
        
        if parser is None:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {list(self.parsers.keys())}"
            )
        
        # Parse document
        result = parser.parse(file_path, **kwargs)
        
        # Add common metadata
        result['file_name'] = file_path.name
        result['file_size'] = file_path.stat().st_size
        
        return result
    
    def parse_batch(self, file_paths: List[str | Path], **kwargs) -> List[Dict[str, any]]:
        """
        Parse multiple documents
        
        Args:
            file_paths: List of file paths
            **kwargs: Additional arguments passed to parser
            
        Returns:
            List of parsed document dictionaries
        """
        results = []
        
        for file_path in file_paths:
            try:
                result = self.parse(file_path, **kwargs)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to parse {file_path}: {e}")
                results.append({
                    'file_path': str(file_path),
                    'error': str(e),
                    'success': False
                })
        
        return results
